# coding=utf-8
"""
@File    : tools.py
@Time    : 2024/2/29 11:21
@Author  : MengYue Sun ldspdvsun@gmail.com
@Description : 基于docx、python-docx和openpyxl的word文档工具类
"""

import datetime
import os
import openpyxl
import string
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx import Document
# 在doc开头调用进行设置
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Mm, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter

def excel_add_sheet(source_excel, sheet_name):
    """
    在 Excel 文件中新增一个工作表。

    :param source_excel: str, Excel 文件路径。
    :param sheet_name: str, 新工作表的名称。
    :return: bool, 添加操作是否成功。
    """
    try:
        # 加载工作簿
        workbook = openpyxl.load_workbook(source_excel)
        # 检查是否已存在相同名称的工作表
        if sheet_name in workbook.sheetnames:
            print(f"Sheet '{sheet_name}' already exist.")
            return False
        # 添加新工作表
        workbook.create_sheet(title=sheet_name)
        # 保存更改
        workbook.save(source_excel)
        print(f"Successful add new sheet: '{sheet_name}'")
        return True
    except FileNotFoundError:
        print("Excel not found")
        return False
    except Exception as e:
        print(f"add new sheet error: {e}")
        return False


def excel_copy_file(source_excel, target_excel):
    """
    复制 Excel 文件到目标路径。如果目标文件已存在，提示并终止复制。
    :param source_excel: 源文件路径。
    :param target_excel: 目标文件路径。
    :return: True 表示复制成功，False 表示失败。
    """
    try:
        # 检查源文件是否存在
        if not os.path.exists(source_excel):
            raise FileNotFoundError(f"Source file '{source_excel}' does not exist.")

        # 检查目标文件是否已存在
        if os.path.exists(target_excel):
            print(f"Target file '{target_excel}' already exists. Aborting copy operation.")
            return False

        # 加载源文件内容
        workbook = openpyxl.load_workbook(source_excel)

        # 保存到目标文件
        workbook.save(target_excel)
        print(f"Excel file copied to: {target_excel}")
        return True
    except Exception as e:
        print(f"An error occurred in copy_file: {e}")
        return False


def excel_count_sheet_columns_with_prefix_suffix(sheet, column_prefix='prefix', column_suffix='suffix'):
    """
    统计符合命名规则的列的值出现次数，并汇总结果。

    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param column_prefix: str, 目标列名称的前缀（默认为 'prefix'）。
    :param column_suffix: str, 目标列名称的后缀（默认为 'suffix'）。
    :return: list[dict], 每个Sheet的统计结果列表，每个字典包含匹配到的目标名称及其执行结果统计。
    """
    try:
        # 获取表头信息（第一行）
        header = [cell.value for cell in sheet[1]]

        # 初始化存储各个平台执行情况的列表
        list_platform_results = []

        # 遍历表头中所有列
        for i, col_name in enumerate(header):
            # 判断列名是否符合前缀和后缀规则
            if isinstance(col_name, str) and col_name.startswith(column_prefix) and col_name.endswith(
                    column_suffix):
                try:
                    # 提取名称
                    platform_name = col_name[len(column_prefix):-len(column_suffix)].strip()

                    # 初始化当前的统计情况字典
                    platform_results = {'target': platform_name, 'results': {}}

                    # 找到当前列的索引
                    platform_column_index = i + 1

                    # 统计当前列的值及其出现次数
                    platform_counts = {}
                    for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, min_col=platform_column_index,
                                               max_col=platform_column_index, values_only=True):
                        result = row[0]  # 获取当前行的执行结果
                        if result is not None:
                            platform_counts[result] = platform_counts.get(result, 0) + 1

                    # 将统计结果保存到当前平台的执行情况字典中
                    platform_results['results'] = platform_counts

                    # 将当前平台的执行情况添加到结果列表中
                    list_platform_results.append(platform_results)

                except Exception as e:
                    print(f"Error processing column '{col_name}': {e}")

        return list_platform_results

    except Exception as e:
        print(f"An error occurred in excel_count_sheet_columns_with_prefix_suffix: {e}")
        return []


def excel_delete_file(target_excel):
    """
    删除指定的 Excel 文件。
    :param target_excel: 目标文件路径。
    :return: True 表示删除成功，False 表示失败。
    """
    try:
        # 检查文件是否存在
        if not os.path.exists(target_excel):
            raise FileNotFoundError(f"File '{target_excel}' does not exist.")

        # 删除文件
        os.remove(target_excel)
        print(f"Excel file deleted: {target_excel}")
        return True
    except Exception as e:
        print(f"An error occurred in delete_file: {e}")
        return False


def excel_delete_sheet_by_identifiers(source_excel, sheet_identifiers):
    """
    从 Excel 文件中删除指定的工作表，支持按索引或名称删除。

    :param source_excel: str, Excel 文件路径。
    :param sheet_identifiers: int 或 str, 工作表索引（从 0 开始）或工作表名称。
    :return: bool, 删除操作是否成功。
    """
    try:
        # 加载工作簿
        workbook = openpyxl.load_workbook(source_excel)
        sheet_names = workbook.sheetnames  # 获取所有工作表名称

        # 如果是单个索引或名称，统一转换为列表处理
        if isinstance(sheet_identifiers, (int, str)):
            sheet_identifiers = [sheet_identifiers]

        # 初始化删除结果和错误列表
        deleted_sheets = []
        errors = []

        for identifier in sheet_identifiers:
            if isinstance(identifier, int):
                # 验证索引有效性
                if 0 <= identifier < len(sheet_names):
                    sheet_name_to_delete = sheet_names[identifier]
                    del workbook[sheet_name_to_delete]
                    deleted_sheets.append(sheet_name_to_delete)
                else:
                    errors.append(
                        f"Invalid sheet index: {identifier}. Must be between 0 and {len(sheet_names) - 1}.")
            elif isinstance(identifier, str):
                # 验证工作表名称是否存在
                if identifier in sheet_names:
                    del workbook[identifier]
                    deleted_sheets.append(identifier)
                else:
                    errors.append(f"Sheet name '{identifier}' not found in workbook.")
            else:
                errors.append(f"Invalid sheet identifier type: {type(identifier).__name__}. Must be int or str.")

        # 保存更改
        workbook.save(source_excel)

        # 打印成功删除的工作表
        for sheet in deleted_sheets:
            print(f"Successfully deleted sheet: '{sheet}'")

        # 打印错误信息
        if errors:
            print("The following errors occurred:")
            for error in errors:
                print(f"  - {error}")

        return not errors  # 如果没有错误，则返回 True，否则返回 False

    except FileNotFoundError:
        print("Excel file not found. ")
        return False
    except Exception as e:
        print(f"Delete error: {e}")
        return False


def excel_delete_sheet_columns(sheet, target_column_index, number_of_columns):
    """
    在指定的 Sheet 工作表中，从目标列索引开始删除指定数量的列。

    :param sheet: openpyxl.Worksheet 对象，指定的工作表，excel_get_sheet()返回值。
    :param target_column_index: int, 要删除的起始列索引，从 1 开始。
    :param number_of_columns: int, 要删除的列数。
    :return: bool, 删除操作是否成功。
    """
    try:
        # 删除指定列
        sheet.delete_cols(target_column_index, number_of_columns)
        print(
            f"Successfully deleted {number_of_columns} column(s) starting from column {target_column_index} in '{sheet.title}'.")
        return True
    except Exception as e:
        print(f"Delete error: {e}")
        return False


def excel_delete_sheet_columns_by_identifiers(sheet, column_identifiers):
    """
    在指定的工作表中，根据列名称或索引删除对应的列（支持单个或多个）。

    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param column_identifiers: str, int, 或 list, 列名称或列索引（从 1 开始），支持单个或多个。
    :return: bool, 删除操作是否成功。
    """
    try:
        # 将单个列名或索引转换为列表，统一处理
        if isinstance(column_identifiers, (str, int)):
            column_identifiers = [column_identifiers]

        # 获取列标题行（假设第一行为标题）
        headers = [cell.value for cell in sheet[1]]

        # 转换列名称为索引，并验证索引的有效性
        columns_to_delete = []
        for identifier in column_identifiers:
            if isinstance(identifier, int):
                # 验证索引范围
                if 1 <= identifier <= sheet.max_column:
                    columns_to_delete.append(identifier)
                else:
                    print(f"Invalid column index: {identifier}. Must be between 1 and {sheet.max_column}.")
            elif isinstance(identifier, str):
                # 检查列名是否存在于标题行
                if identifier in headers:
                    col_index = headers.index(identifier) + 1  # 列索引从 1 开始
                    columns_to_delete.append(col_index)
                else:
                    print(f"Column name '{identifier}' not found in sheet '{sheet.title}'.")
            else:
                print(f"Invalid column identifier type: {type(identifier).__name__}. Must be int or str.")

        # 如果没有找到要删除的列
        if not columns_to_delete:
            print("No valid columns found to delete. No changes made.")
            return False

        # 按降序删除列，防止索引混乱
        columns_to_delete.sort(reverse=True)

        for col_index in columns_to_delete:
            sheet.delete_cols(col_index)

        print(f"Successfully deleted columns: {column_identifierss} from sheet '{sheet.title}'.")
        return True

    except Exception as e:
        print(f"An error occurred while deleting columns: {e}")
        return False


def excel_delete_sheet_rows(sheet, target_row_index, number_of_rows):
    """
    在指定的 Sheet 工作表中，从目标行索引开始删除指定数量的行。

    :param sheet: openpyxl.Worksheet 对象，指定的工作表，excel_get_sheet()返回值。
    :param ttarget_row_index: int, 要删除的起始行索引，从 1 开始。
    :param number_of_rows: int, 要删除的行数。
    :return: bool, 删除操作是否成功。
    """
    try:
        # 删除指定行
        sheet.delete_rows(target_row_index + 1, number_of_rows)
        print(
            f"Successfully deleted {number_of_rows} column(s) starting from column {target_row_index} in '{sheet.title}'.")
        return True
    except Exception as e:
        print(f"Delete error: {e}")
        return False


def excel_fill_sheet_cells(sheet, ranges_to_fill, fill_value=None):
    """
    填充 Excel 中指定的单元格区域。
    支持单一值填充或按区域逐一填充。

    :param sheet: openpyxl.Worksheet 对象，指定的工作表，excel_get_sheet()返回值。
    :param ranges_to_fill: tuple 或 list[tuple], 要填充的单个区域或区域列表。
                           单个区域格式为 (start_row, start_col, end_row, end_col)；
                           多个区域格式为 [(start_row, start_col, end_row, end_col), ...]。
    :param fill_value: 任意, 填充的值。
                       - 若为 None，则跳过值填充（用于逐单元格自定义填充）。
                       - 若为列表或二维列表，需与区域大小匹配。
    :return: None
    """
    try:
        # 将单个区域统一处理为列表格式
        if isinstance(ranges_to_fill, tuple):
            ranges_to_fill = [ranges_to_fill]

        # 遍历所有区域
        for cell_range in ranges_to_fill:
            start_row, start_col, end_row, end_col = cell_range
            for row in range(start_row, end_row + 1):
                for col in range(start_col, end_col + 1):
                    cell = sheet.cell(row=row, column=col)
                    if fill_value is not None:
                        # 检查填充值是否为二维列表
                        if isinstance(fill_value, list) and isinstance(fill_value[0], list):
                            cell.value = fill_value[row - start_row][col - start_col]
                        # 检查填充值是否为一维列表
                        elif isinstance(fill_value, list):
                            cell.value = fill_value[row - start_row]
                        else:
                            # 单一值填充
                            cell.value = fill_value
        print(f"Specified ranges filled successfully in sheet '{sheet.title}'.")

    except Exception as e:
        print(f"An error occurred in excel_set_sheet_cells_value: {e}")


def excel_get_sheet(workbook, sheet_identifiers):
    """
    根据工作表索引或名称获取工作表对象。

    :param workbook: openpyxl.Workbook, 要操作的工作簿对象。
    :param sheet_identifiers: int 或 str, 工作表索引（从 0 开始）或工作表名称。
    :return: openpyxl.Worksheet, 对应的工作表对象。
    :raises ValueError: 如果工作表索引或名称无效。
    """
    try:
        # 获取所有的工作表名称
        sheet_names = workbook.sheetnames
        # 如果是整数，按索引获取工作表
        if isinstance(sheet_identifiers, int):
            if 0 <= sheet_identifiers < len(sheet_names):
                return workbook.worksheets[sheet_identifiers]
            else:
                raise ValueError(
                    f"Invalid sheet index: {sheet_identifiers}. Must be between 0 and {len(workbook.worksheets) - 1}.")

        # 如果是字符串，按名称获取工作表
        elif isinstance(sheet_identifiers, str):
            if sheet_identifiers in workbook.sheetnames:
                return workbook[sheet_identifiers]
            else:
                raise ValueError(f"Sheet name '{sheet_identifiers}' not found in workbook.")

        # 非法的 sheet_identifiers 类型
        else:
            raise ValueError("sheet_identifiers must be an integer (index) or a string (sheet name).")
    except Exception as e:
        raise RuntimeError(f"An error occurred while retrieving the sheet: {e}")


def excel_get_sheet_cell_date_value(sheet, row, column):
    """
    获取指定单元格的日期值，并将其格式化为字符串返回。

    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param row: int, 行索引。
    :param column: int, 列索引。
    :return: str, 单元格的日期值的字符串表示形式（格式为'%Y年%m月%d日'）
    """
    try:
        date_value = sheet.cell(row, column).value
        if isinstance(date_value, datetime.datetime):
            return date_value.strftime('%Y年%m月%d日')
        else:
            return str(date_value)
    except Exception as e:
        print(f"An error occurred in get_sheet_cell_date_value: {e}")
        return ""


def excel_get_sheet_cell_value(sheet, row, column):
    """
    获取指定单元格的值，并将其转换为字符串返回。

    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param row: int, 行索引。
    :param column: int, 列索引。
    :return: str, 单元格的值的字符串表示形式。
    """
    try:
        return str(sheet.cell(row, column).value)
    except Exception as e:
        print(f"An error occurred in get_sheet_cell_value: {e}")
        return ""


def excel_get_sheet_column_keyword_count(sheet, column_identifiers, keyword):
    """
    遍历指定工作表的某列，并返回包含指定关键词的单元格数量。
    支持按列名或列索引指定列。

    :param sheet: openpyxl.Worksheet, 工作表对象。
    :param column_identifiers: str 或 int, 列名（标题）或列索引（从 1 开始）。
    :param keyword: str, 要统计的关键词。
    :return: int, 包含指定关键词的单元格数量。如果发生错误，返回 -1。
    """
    try:
        # 确定列索引
        if isinstance(column_identifiers, int):
            # 验证列索引范围
            if column_identifiers < 1 or column_identifiers > sheet.max_column:
                raise ValueError(
                    f"Invalid column index: {column_identifiers}. Must be between 1 and {sheet.max_column}.")
            column_index = column_identifiers
        elif isinstance(column_identifiers, str):
            # 获取列标题所在行
            headers = [cell.value for cell in sheet[1]]
            if column_identifier not in headers:
                raise ValueError(f"Column '{column_identifier}' not found in sheet '{sheet.title}'.")
            column_index = headers.index(column_identifier) + 1
        else:
            raise ValueError("Column identifier must be an integer (index) or a string (name).")

        # 统计包含指定关键词的单元格数量
        keyword_count = sum(
            1 for row in
            sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=column_index, max_col=column_index)
            if row[0].value is not None and str(row[0].value) == keyword
        )

        return keyword_count

    except Exception as e:
        print(f"An error occurred in excel_get_column_keyword_count: {e}")
        return -1  # 返回 -1 表示发生错误


def excel_get_sheet_column_len(sheet, column_identifiers):
    """
    遍历指定工作表的某列，返回该列的非重复值长度。
    支持按列索引或列名称指定列。

    :param sheet: openpyxl.Worksheet, 工作表对象。
    :param column_identifiers: int 或 str, 列索引（从 1 开始）或列名称。
    :return: int, 非重复值的列长度。如果发生错误，返回 -1。
    """
    try:
        # 确定列索引
        if isinstance(column_identifiers, int):
            # 验证列索引范围
            if column_identifiers < 1 or column_identifiers > sheet.max_column:
                raise ValueError(
                    f"Invalid column index: {column_identifiers}. Must be between 1 and {sheet.max_column}.")
            column_index = column_identifiers
        elif isinstance(column_identifiers, str):
            # 获取列标题所在行
            headers = [cell.value for cell in sheet[1]]
            if column_identifiers not in headers:
                raise ValueError(f"Column '{column_identifiers}' not found in sheet '{sheet.title}'.")
            column_index = headers.index(column_identifiers) + 1
        else:
            raise ValueError("Column identifiers must be an integer (index) or a string (name).")

        # 使用集合来跟踪已经出现过的值
        unique_values = set()

        # 遍历指定列的所有数据，从第1行开始
        for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=column_index, max_col=column_index):
            value = row[0].value
            if value is not None:
                unique_values.add(value)

        # 返回非重复值的长度
        return len(unique_values)

    except Exception as e:
        print(f"An error occurred in excel_get_column_len: {e}")
        return -1  # 返回 -1 表示发生错误


def excel_get_sheet_column_str_custom(sheet, column_identifiers, separator='、', end_mark='。'):
    """
    遍历指定工作表的某列，并返回该列的所有数据拼接成的字符串。
    支持按列索引或列名称指定列。

    :param sheet: openpyxl.Worksheet, 工作表对象。
    :param column_identifiers: int 或 str, 列索引（从 1 开始）或列名称。
    :param separator: str, 单元格值之间的分隔符，默认为 '、'。
    :param end_mark: str, 列字符串的结束标记，默认为 '。'。
    :return: str, 拼接后的字符串。如果发生错误，返回空字符串。
    """
    try:
        # 确定列索引
        if isinstance(column_identifiers, int):
            # 验证列索引范围
            if column_identifiers < 1 or column_identifiers > sheet.max_column:
                raise ValueError(
                    f"Invalid column index: {column_identifiers}. Must be between 1 and {sheet.max_column}.")
            column_index = column_identifiers
        elif isinstance(column_identifiers, str):
            # 获取列标题所在行
            headers = [cell.value for cell in sheet[1]]
            if column_identifiers not in headers:
                raise ValueError(f"Column '{column_identifiers}' not found in sheet '{sheet.title}'.")
            column_index = headers.index(column_identifiers) + 1
        else:
            raise ValueError("Column identifiers must be an integer (index) or a string (name).")

        # 存储列数据的列表
        column_data = []

        # 遍历指定列的所有数据，从第二行开始
        for row in range(2, sheet.max_row + 1):
            cell_value = sheet.cell(row=row, column=column_index).value
            # 检查单元格的值是否为 None，如果是则跳过
            if cell_value is not None:
                column_data.append(str(cell_value))

        # 拼接列数据为字符串
        column_string = separator.join(column_data)

        # 在字符串末尾添加结束标记（如果数据不为空）
        if column_string:
            column_string += end_mark

        return column_string

    except Exception as e:
        print(f"An error occurred in excel_get_column_str_custom: {e}")
        return ""  # 返回空字符串表示发生了错误


def excel_get_sheet_column_values(sheet, column_identifiers, include_header=True):
    """
    从指定工作表中获取列的所有值，支持按列索引或列名称获取。

    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param column_identifiers: int 或 str 或 list, 列索引（从 1 开始）或列名称（单个或多个）。
    :param include_header: bool, 是否包含标题行（默认为 True）。
    :return: dict 或 list, 单个列返回列表，多个列返回字典（仅包含有效列）。
    """
    try:
        # 如果是单个列名或索引，统一转换为列表处理
        if isinstance(column_identifiers, (str, int)):
            column_identifiers = [column_identifiers]

        # 获取标题行（假设第一行为标题）
        headers = [cell.value for cell in sheet[1]]

        # 初始化结果和错误列表
        column_indices = {}
        errors = []

        for identifier in column_identifiers:
            if isinstance(identifier, int):
                # 验证列索引是否有效
                if 1 <= identifier <= sheet.max_column:
                    column_name = headers[identifier - 1] if identifier <= len(headers) else f"Column {identifier}"
                    column_indices[column_name] = identifier
                else:
                    errors.append(f"Invalid column index: {identifier}. Must be between 1 and {sheet.max_column}.")
            elif isinstance(identifier, str):
                # 验证列名称是否存在
                if identifier in headers:
                    column_indices[identifier] = headers.index(identifier) + 1
                else:
                    errors.append(f"Column name '{identifier}' not found in sheet '{sheet.title}'.")
            else:
                errors.append(f"Invalid column identifier type: {type(identifier).__name__}. Must be int or str.")

        # 提取有效列的值
        result = {}
        for col_name, col_index in column_indices.items():
            values = [sheet.cell(row=row, column=col_index).value for row in range(1, sheet.max_row + 1)]
            if not include_header:
                values = values[1:]  # 去掉标题行
            result[col_name] = values

        # 打印无效列的错误信息
        if errors:
            print("The following errors occurred:")
            for error in errors:
                print(f"  - {error}")

        # 如果只有一个有效列，直接返回列表
        if len(result) == 1:
            return list(result.values())[0]

        return result

    except Exception as e:
        raise RuntimeError(f"An unexpected error occurred while retrieving column values: {e}")


def excel_get_sheet_row_col_count(sheet):
    """
    获取指定工作表的总行数和总列数。

    :param sheet: openpyxl.Worksheet 对象，指定的工作表，excel_get_sheet()返回值。
    :return: (总行数, 总列数)。
    :raises ValueError: 如果工作表索引或名称无效。
    """
    try:
        # 检查 sheet 是否为有效的工作表对象
        if sheet is None:
            raise ValueError("Invalid sheet object provided.")
        # 获取行数和列数
        total_rows = sheet.max_row
        total_columns = sheet.max_column
        return total_rows, total_columns
    except Exception as e:
        raise RuntimeError(f"An error occurred while retrieving the sheet: {e}")


def excel_get_sheet_row_or_col_max(sheet, identifiers, mode='col'):
    """
    获取指定行的最大列索引，或指定列的最大行索引。

    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param identifiers: int, str 或 list[int | str], 指定的行号/列名或其列表。
    :param mode: str, 'row' 表示获取指定行的最大列，'col' 表示获取指定列的最大行。
    :return: list, 包含元组或字典的列表，每个元素对应指定行或列的最大索引信息。
    """
    try:
        # 将单个输入统一处理为列表
        if isinstance(identifiers, (int, str)):
            identifiers = [identifiers]

        results = []  # 存储结果

        if mode == 'row':
            for identifier in identifiers:
                if isinstance(identifier, int):  # 数字行号
                    row_index = identifier
                elif isinstance(identifier, str):  # 字符串转换为行号
                    row_index = int(identifier)
                else:
                    raise ValueError(f"Invalid row identifier type: {type(identifier)}")

                # 获取该行的最大列索引（最后一个非空单元格）
                max_col = max((col for col in range(1, sheet.max_column + 1)
                               if sheet.cell(row=row_index, column=col).value is not None),
                              default=0)
                results.append({'row': row_index, 'max_col': max_col})

        elif mode == 'col':
            for identifier in identifiers:
                if isinstance(identifier, int):  # 数字列号
                    col_index = identifier
                    col_letter = openpyxl.utils.get_column_letter(col_index)
                elif isinstance(identifier, str):  # 字符列名转换为索引
                    col_index = openpyxl.utils.column_index_from_string(identifier)
                    col_letter = identifier
                else:
                    raise ValueError(f"Invalid column identifier type: {type(identifier)}")

                # 获取该列的最大行索引（最后一个非空单元格）
                max_row = max((row for row in range(1, sheet.max_row + 1)
                               if sheet.cell(row=row, column=col_index).value is not None),
                              default=0)
                results.append({'col': col_letter, 'max_row': max_row})

        else:
            raise ValueError("Invalid mode. Must be 'row' or 'col'.")

        return results

    except Exception as e:
        print(f"An error occurred while retrieving max row or column: {e}")
        return []


def excel_get_sheet_row_values(sheet, row_identifiers, include_index=True):
    """
    从指定工作表中获取行的所有值，支持按行索引获取单行或多行数据。

    :param sheet: openpyxl.Worksheet 对象，指定的工作表，excel_get_sheet()返回值。
    :param row_identifiers: int 或 list[int], 单个行索引或多个行索引，从 1 开始。
    :param include_index: bool, 是否包含行索引（默认为 True，返回字典带行号）。
    :return: dict 或 list, 单行返回列表，多个行返回字典（行索引为键，行值为列表）。
    """
    try:
        # 如果是单个行索引，统一转换为列表处理
        if isinstance(row_identifiers, int):
            row_identifiers = [row_identifiers]

        # 初始化结果和错误列表
        result = {}
        errors = []

        for row_index in row_identifiers:
            # 验证行索引是否有效
            if 1 <= row_index <= sheet.max_row:
                # 获取该行所有单元格的值
                row_values = [sheet.cell(row=row_index, column=col).value for col in range(1, sheet.max_column + 1)]
                result[row_index] = row_values
            else:
                errors.append(f"Invalid row index: {row_index}. Must be between 1 and {sheet.max_row}.")

        # 打印无效行的错误信息
        if errors:
            print("The following errors occurred:")
            for error in errors:
                print(f"  - {error}")

        # 如果只有一个有效行，直接返回列表
        if len(result) == 1:
            return list(result.values())[0] if include_index else list(result.values())[0]

        # 如果需要去掉索引，仅返回行值
        if not include_index:
            return list(result.values())

        return result

    except Exception as e:
        raise RuntimeError(f"An unexpected error occurred while retrieving row values: {e}")


def excel_get_sheets_index_name(workbook):
    """
    获取当前 Workbook 中所有工作表的索引及名称。

    :param workbook: openpyxl.Workbook, 要操作的工作簿对象。
    :return: list[tuple], 每个元素是 (sheet_index, sheet_name) 的元组。
    """
    try:
        if not workbook:
            raise ValueError("Invalid Workbook object. Please provide a valid Workbook.")

        # 获取所有工作表的名称
        sheet_names = workbook.sheetnames
        # 构造索引和名称的列表
        sheet_info = [(index, name) for index, name in enumerate(sheet_names)]

        return sheet_info
    except Exception as e:
        raise RuntimeError(f"An error occurred while retrieving sheet information: {e}")


def excel_get_sheets_targets_counts(list_sheets, column_prefix='prefix', column_suffix='suffix'):
    # def excel_get_keyword_counts_across_sheets(list_sheets, column_prefix='prefix', column_suffix='suffix'):
    """
    获取多个工作表中符合命名规则（前缀和后缀）的列的统计情况。

    :param list_sheets: list[openpyxl.Worksheet], 工作表对象列表。
    :param column_prefix: str, 列名前缀（默认为 'prefix'）。
    :param column_suffix: str, 列名后缀（默认为 'suffix'）。
    :return: list[dict], 每个字典包含工作表名称和列统计结果。
    """
    try:
        # 用于存储所有工作表的统计结果
        combined_results = []

        # 遍历所有工作表
        for sheet in list_sheets:
            sheet_name = sheet.title

            # 调用统计方法获取当前工作表的列统计结果
            sheet_results = excel_count_sheet_columns_with_prefix_suffix(
                sheet,
                column_prefix=column_prefix,
                column_suffix=column_suffix
            )

            # 为当前工作表的每个统计结果添加工作表名称
            for result in sheet_results:
                result['sheet_name'] = sheet_name

            # 合并到总结果列表
            combined_results.extend(sheet_results)

        return combined_results

    except Exception as e:
        print(f"An error occurred in excel_get_keyword_counts_across_sheets: {e}")
        return []


def excel_get_workbook(source_excel):
    """
    加载指定的 Excel 文件并返回 workbook 对象，同时记录原始路径;
    如果作出修改, 需要在结束前调用excel_save_workbook(workbook, output_file=None)保存对象。
    :param source_excel: str, Excel 文件路径。
    :return: openpyxl.Workbook 对象。
    """
    try:
        if not source_excel or not isinstance(source_excel, str):
            raise ValueError("Invalid file path. Please provide a valid Excel file path.")

        # 加载工作簿
        workbook = openpyxl.load_workbook(source_excel)
        # 在 workbook 上记录原始路径
        workbook._original_path = source_excel
        return workbook

    except FileNotFoundError:
        raise FileNotFoundError(f"File not found: {source_excel}. Please check the file path.")
    except Exception as e:
        raise RuntimeError(f"An unexpected error occurred while opening the Excel file: {e}")


def excel_insert_sheet_columns(sheet, target_column_index, number_of_columns):
    """
    在指定工作表中的目标列后插入指定数量的列。

    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param target_column_index: int, 目标列索引，从 1 开始。
    :param number_of_columns: int, 需要插入的列数。
    :return: bool, 插入操作是否成功。
    """
    try:
        # 验证列索引有效性
        if target_column_index < 1 or target_column_index > sheet.max_column + 1:
            raise ValueError(
                f"Invalid target column index: {target_column_index}. Must be between 1 and {sheet.max_column + 1}.")

        # 在目标列后插入列（目标列索引 + 1）
        insert_column_index = target_column_index + 1
        sheet.insert_cols(insert_column_index, number_of_columns)
        print(
            f"Successfully inserted {number_of_columns} column(s) after column {target_column_index} in sheet '{sheet.title}'.")
        return True

    except Exception as e:
        print(f"An error occurred while inserting columns: {e}")
        return False


def excel_insert_sheet_rows(sheet, target_row_index, number_of_rows):
    """
    在指定工作表中的目标列后插入指定数量的行。

    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param target_row_index: int, 目标列索引，从 1 开始。
    :param number_of_rows: int, 需要插入的行数。
    :return: bool, 插入操作是否成功。
    """
    try:
        # 验证列索引有效性
        if target_row_index < 1 or target_row_index > sheet.max_row + 1:
            raise ValueError(
                f"Invalid target column index: {target_column_index}. Must be between 1 and {sheet.max_column + 1}.")

            # 在目标行的下一行插入行
            insert_row_index = target_row_index + 1
            sheet.insert_rows(insert_row_index, number_of_rows)
            print(
                f"Successfully inserted {number_of_rows} row(s) after row {target_row_index} in sheet '{sheet.title}'.")
            return True
    except Exception as e:
        print(f"An error occurred while inserting rows: {e}")
        return False


def excel_merge_sheet_cells(sheet, ranges_to_merge):
    """
    合并指定工作表中的单元格区域，并设置内容居中对齐（上下、左右）。
    默认使用合并区域左上角单元格的值作为合并后单元格的值。

    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param ranges_to_merge: tuple 或 list[tuple], 要合并的单个区域或区域列表。
                            单个区域格式为 (start_row, start_col, end_row, end_col)；
                            多个区域格式为 [(start_row, start_col, end_row, end_col), ...]。
    :return: None
    """
    try:
        # 将单个区域统一处理为列表格式
        if isinstance(ranges_to_merge, tuple):
            ranges_to_merge = [ranges_to_merge]

        # 遍历所有要合并的区域
        for cell_range in ranges_to_merge:
            start_row, start_col, end_row, end_col = cell_range

            # 合并单元格
            sheet.merge_cells(start_row=start_row, start_column=start_col, end_row=end_row, end_column=end_col)

            # 获取合并区域左上角单元格
            cell = sheet.cell(row=start_row, column=start_col)

            # 设置合并单元格的内容居中对齐
            cell.alignment = Alignment(horizontal="center", vertical="center")

        print(f"Successfully merged specified ranges in sheet '{sheet.title}'.")

    except Exception as e:
        print(f"An error occurred while merging cells: {e}")


def excel_replace_sheet_content(sheet, target_text, replacement_text, rows=None, cols=None):
    """
    在 Excel 文件中查找并替换内容，支持遍历所有单元格，也支持指定行或列。
    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param target_text: 要查找的文本。
    :param replacement_text: 替换的文本。
    :param rows: 指定的行列表（如 1或[2,3]）从 1 开始；如果为 None，则遍历所有行。
    :param cols: 指定的列列表（如 1或[2,3]）从 1 开始；如果为 None，则遍历所有列。
    """
    try:

        # 处理 rows 和 cols 参数，确保它们是可迭代对象
        if isinstance(rows, int):
            rows = [rows]  # 将单个行号包装为列表
        if isinstance(cols, int):
            cols = [cols]  # 将单个列号包装为列表

        # 获取行范围和列范围
        row_range = rows or range(1, sheet.max_row + 1)  # 遍历所有行或指定行
        col_range = cols or range(1, sheet.max_column + 1)  # 遍历所有列或指定列

        for row in row_range:
            for col in col_range:
                cell = sheet.cell(row=row, column=col)
                if cell.value and isinstance(cell.value, str) and target_text in cell.value:
                    # 替换内容
                    new_value = cell.value.replace(target_text, replacement_text)
                    cell.value = new_value
                    modified = True
    except Exception as e:
        print(f"An error occurred in excel_sheet_replace_content: {e}")


def excel_save_workbook(workbook, output_file=None):
    """
    保存 Workbook 对象到指定路径。如果未指定路径，则保存到加载时的原路径。
    :param workbook: openpyxl.Workbook, 要保存的工作簿对象。
    :param output_file: str, 输出文件路径。如果为 None，则保存到加载时记录的路径。
    """
    try:
        if not workbook:
            raise ValueError("Invalid Workbook object. Cannot save an empty workbook.")

        # 如果未指定输出路径，则尝试使用 Workbook 中记录的原路径
        if not output_file:
            if hasattr(workbook, "_original_path") and workbook._original_path:
                save_path = workbook._original_path
            else:
                raise ValueError("No save path provided, and workbook has no recorded original path.")
        else:
            save_path = output_file

        # 保存工作簿
        workbook.save(save_path)
        print(f"Workbook saved successfully to: {save_path}")
    except Exception as e:
        raise RuntimeError(f"An error occurred while saving the workbook: {e}")


def excel_set_sheet_alignment(sheet, ranges, horizontal="center", vertical="center"):
    """
    设置指定区域的单元格对齐方式。
    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param sheet_index: 工作表索引（从 0 开始）。
    :param ranges: 单个区域 (start_row, start_col, end_row, end_col) 或多个区域 [(..), (..)]。
    :param horizontal: 水平对齐方式（如 "left", "center", "right"）。
    :param vertical: 垂直对齐方式（如 "top", "center", "bottom"）。
    """
    try:
        if isinstance(ranges, tuple):
            ranges = [ranges]

        for cell_range in ranges:
            start_row, start_col, end_row, end_col = cell_range
            for row in range(start_row, end_row + 1):
                for col in range(start_col, end_col + 1):
                    cell = sheet.cell(row=row, column=col)
                    cell.alignment = openpyxl.styles.Alignment(horizontal=horizontal, vertical=vertical)
    except Exception as e:
        print(f"An error occurred in excel_set_alignment: {e}")


def excel_set_sheet_bold(sheet, ranges, bold=True):
    """
    设置指定区域的字体加粗。
    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param ranges: 单个区域 (start_row, start_col, end_row, end_col) 或多个区域 [(..), (..)]。
    :param bold: 是否加粗（True/False）。
    """
    try:
        if isinstance(ranges, tuple):
            ranges = [ranges]

        for cell_range in ranges:
            start_row, start_col, end_row, end_col = cell_range
            for row in range(start_row, end_row + 1):
                for col in range(start_col, end_col + 1):
                    cell = sheet.cell(row=row, column=col)
                    if cell.value is not None:
                        cell.font = openpyxl.styles.Font(bold=bold)
    except Exception as e:
        print(f"An error occurred in excel_sheet_set_bold: {e}")


def excel_set_sheet_cell_value(sheet, row, column, value):
    """
    在指定的工作表的指定行列设置值。

    :param sheet: openpyxl.Worksheet, 工作表对象。
    :param row: int, 目标单元格所在行，从 1 开始。
    :param column: int, 目标单元格所在列，从 1 开始。
    :param value: 任意, 要设置的值。
    :return: bool, 设置操作是否成功。
    """
    try:
        # 验证行和列索引是否有效
        if row < 1 or column < 1:
            raise ValueError("Row and column indices must be greater than or equal to 1.")

        # 设置指定单元格的值
        sheet.cell(row=row, column=column, value=value)
        print(f"Value set successfully in sheet '{sheet.title}' at row {row}, column {column}: {value}")
        return True

    except Exception as e:
        print(f"Error while setting value: {e}")
        return False


def excel_set_sheet_fill_color(sheet, ranges, color_hex):
    """
    设置指定区域的单元格填充颜色。
    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param ranges: 单个区域 (start_row, start_col, end_row, end_col) 或多个区域 [(..), (..)]。
    :param color_hex: 填充颜色（十六进制字符串，如 "FFFF00" 表示黄色）。
    """
    try:
        if isinstance(ranges, tuple):
            ranges = [ranges]

        fill = openpyxl.styles.PatternFill(start_color=color_hex, end_color=color_hex, fill_type="solid")

        for cell_range in ranges:
            start_row, start_col, end_row, end_col = cell_range
            for row in range(start_row, end_row + 1):
                for col in range(start_col, end_col + 1):
                    cell = sheet.cell(row=row, column=col)
                    cell.fill = fill
    except Exception as e:
        print(f"An error occurred in excel_sheet_set_fill_color: {e}")


def excel_set_sheet_font_color(sheet, ranges, color_hex):
    """
    设置指定区域的字体颜色。
    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param ranges: 单个区域 (start_row, start_col, end_row, end_col) 或多个区域 [(..), (..)]。
    :param color_hex: 字体颜色（十六进制字符串，如 "FF0000" 表示红色）。
    """
    try:
        if isinstance(ranges, tuple):
            ranges = [ranges]

        for cell_range in ranges:
            start_row, start_col, end_row, end_col = cell_range
            for row in range(start_row, end_row + 1):
                for col in range(start_col, end_col + 1):
                    cell = sheet.cell(row=row, column=col)
                    if cell.value is not None:
                        cell.font = openpyxl.styles.Font(color=color_hex)
    except Exception as e:
        print(f"An error occurred in excel_sheet_set_font_color: {e}")


def excel_set_sheet_font_size(sheet, ranges, font_size):
    """
    设置指定区域的字体字号。
    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param ranges: 单个区域 (start_row, start_col, end_row, end_col) 或多个区域 [(..), (..)]。
    :param font_size: 字号大小（int）。
    """
    try:
        if isinstance(ranges, tuple):
            ranges = [ranges]
        if not isinstance(font_size, int):
            raise TypeError("font_size type error")

        for cell_range in ranges:
            start_row, start_col, end_row, end_col = cell_range
            for row in range(start_row, end_row + 1):
                for col in range(start_col, end_col + 1):
                    cell = sheet.cell(row=row, column=col)
                    if cell.value is not None:
                        cell.font = openpyxl.styles.Font(size=font_size)
    except Exception as e:
        print(f"An error occurred in excel_sheet_set_font_size: {e}")


def excel_set_sheet_freeze_panes(sheet, cell_reference):
    """
    冻结窗口。
    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param sheet_index: 工作表索引（从 0 开始）。
    :param cell_reference: 冻结的参考单元格（如 "C2" 冻结A、B列, 第1行），必须传递。
    """
    try:
        if not cell_reference:
            raise ValueError("The 'cell_reference' parameter is required and cannot be empty.")

        # 设置冻结窗格
        sheet.freeze_panes = cell_reference
    except Exception as e:
        print(f"An error occurred in excel_sheet_freeze_panes: {e}")


def excel_unmerge_sheet_cells(sheet, ranges_to_unmerge):
    """
    取消指定工作表中的合并单元格，并填充原合并单元格的内容到拆分后的所有单元格中。

    :param sheet: openpyxl.Worksheet, 要操作的工作表对象。
    :param ranges_to_unmerge: tuple 或 list[tuple], 要取消合并的单个区域或区域列表。
                              单个区域格式为 (start_row, start_col, end_row, end_col)；
                              多个区域格式为 [(start_row, start_col, end_row, end_col), ...]。
    :return: None
    """
    try:
        # 将单个区域统一处理为列表格式
        if isinstance(ranges_to_unmerge, tuple):
            ranges_to_unmerge = [ranges_to_unmerge]

        for cell_range in ranges_to_unmerge:
            start_row, start_col, end_row, end_col = cell_range
            range_str = f"{get_column_letter(start_col)}{start_row}:{get_column_letter(end_col)}{end_row}"

            # 检查是否为合并单元格
            if range_str in sheet.merged_cells:
                # 获取合并区域左上角单元格的值
                content = sheet.cell(row=start_row, column=start_col).value

                # 取消合并
                sheet.unmerge_cells(range_str)

                # 填充拆分后的单元格
                for row in range(start_row, end_row + 1):
                    for col in range(start_col, end_col + 1):
                        cell = sheet.cell(row=row, column=col)
                        cell.value = content
                        cell.alignment = Alignment(horizontal="left", vertical="center")
            else:
                print(f"Range {range_str} is not a merged cell. Skipping.")

        print(f"Successfully unmerged specified ranges in sheet '{sheet.title}'.")

    except Exception as e:
        print(f"An error occurred while unmerging cells: {e}")


def word_add_custom_heading(document, level, text):
    """
    向文档添加自定义级别的标题。
    :param document: 文档对象。
    :param level: 标题级别，整数值，1 表示一级标题，2 表示二级标题，以此类推。
    :param text: 标题文本。
    """
    try:
        if level < 1 or level > 9:
            raise ValueError("Heading level must be between 1 and 9")
        else:
            heading_style = f"Heading {level}"
            document.add_heading(text, level=level).style = heading_style

    except ValueError as ve:
        print(f"ValueError in add_custom_heading: {ve}")


def word_add_custom_paragraph(document, text, style=None):
    """
    向文档中添加自定义样式的段落。
    :param document: Document对象，表示要添加段落的文档。
    :param text: 要添加的文本内容。
    :param style: 要应用的样式名称，默认为None。
    """
    try:
        if style:
            document.add_paragraph(text, style=style)
        else:
            document.add_paragraph(text)

    except Exception as e:
        print(f"An error occurred in add_custom_paragraph: {e}")


def word_add_custom_table(document, row_count, column_count, style="Table Grid", font_size=10.5, spacing_before=0):
    """
    向文档中添加自定义样式的表格。
    :param document: Document对象，表示要添加表格的文档。
    :param row_count: 表格的行数。
    :param column_count: 表格的列数。
    :param style: 要应用的表格样式名称，默认为 "Table Grid"。
    :param font_size: 字体大小，默认为10.5。
    :param spacing_before: 段前间距，默认为0。
    """
    try:
        table = document.add_table(rows=row_count, cols=column_count)
        if style:
            table.style = style
        else:
            table.style = "Table Grid"
        # table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)  # 设置字体大小
                    paragraph.paragraph_format.space_before = Pt(spacing_before)  # 设置段前间距
        return table

    except Exception as e:
        print(f"An error occurred in add_custom_table: {e}")


def word_add_figure_caption(document, caption_text):
    """
    在 Word 文档中添加图序号题注
    :param document: Word 文档对象
    :param caption_text: 题注文本
    """
    try:
        # 获取所有表格
        tables = document.tables

        # 获取表格数量
        table_count = len(tables)

        # 获取最近的标题（章节号）
        section_number = None
        for paragraph in document.paragraphs[::-1]:  # 倒序查找
            if paragraph.style.name.startswith('Heading'):
                section_text = get_section_number(paragraph.text)
                if section_text is not None:
                    section_number = section_text.split('.')[0]
                break

        # 创建一个新的段落对象用于存放图序号题注
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐

        # 创建一个新的图序号题注对象
        caption = OxmlElement('w:fldSimple')
        caption.set(qn('w:instr'), f'SEQ Figure \\* ARABIC')
        run = p.add_run()
        run.text = f"图 {section_number}-{table_count + 1}：{caption_text}"
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        p._p.append(caption)

    except Exception as e:
        print(f"An error occurred in add_figure_caption: {e}")


def word_add_indented_paragraph(document, text):
    """
    在 Word 文档中添加一个带有序号和缩进的段落。
    :param document: Word 文档对象
    :param text: 段落文本
    """
    try:
        # 将文本拆分成行
        lines = text.split("\n")

        # 添加段落，并在每一行前添加序号和制表符缩进
        for i, line in enumerate(lines, start=1):
            indented_line = f"\t{i}.{line}"
            document.add_paragraph(indented_line, style="BodyText")

    except Exception as e:
        print(f"An error occurred in add_indented_paragraph: {e}")


def word_add_ordered_list(document, items, index_style='number', indentation=1):
    """
    向文档中添加有序列表，并设置缩进。
    :param document: 文档对象。
    :param items: 有序列表项的列表。
    :param index_style: 索引样式，可以是 'number'（数字）、'lower_alpha'（小写字母）、'upper_alpha'（大写字母）、'roman'（罗马数字）之一，默认为 'number'。
    :param indentation: 缩进量，默认为1级0.5英寸。
    """
    try:
        indexes = []  # 存储索引列表
        if index_style == 'number':
            indexes = [str(i) for i in range(1, len(items) + 1)]
        elif index_style == 'lower_alpha':
            indexes = list(string.ascii_lowercase)[:len(items)]  # 将数字转换为小写字母
        elif index_style == 'upper_alpha':
            indexes = [string.ascii_lowercase[i].upper() for i in range(len(items))]  # 将数字转换为大写字母
        elif index_style == 'roman':
            indexes = [set_number_to_roman(i) for i in range(1, len(items) + 1)]  # 使用 to_roman 方法转换数字为罗马数字
        else:
            raise ValueError("Invalid index style. Use 'number', 'low_alpha', 'upper_alpha', or 'roman'.")

        for index, item in zip(indexes, items):
            # 如果不是最后一项，在末尾添加分号
            if index != indexes[-1]:
                paragraph = document.add_paragraph(f"{index}、{item}；")
            else:
                # 如果是最后一项，在末尾添加句号
                paragraph = document.add_paragraph(f"{index}、{item}。")
            paragraph.paragraph_format.left_indent = Inches(indentation * 0.5)

    except Exception as e:
        print(f"An error occurred in add_ordered_list: {e}")


def word_add_table_caption(document, caption_text):
    """
    在 Word 文档中添加表序号题注
    :param document: Word 文档对象
    :param caption_text: 题注文本
    """
    try:
        # 获取所有表格
        tables = document.tables

        # 获取表格数量
        table_count = len(tables)

        # 获取最近的标题（章节号）
        section_number = None
        for paragraph in document.paragraphs[::-1]:  # 倒序查找
            if paragraph.style.name.startswith('Heading'):
                section_text = get_section_number(paragraph.text)
                if section_text is not None:
                    section_number = section_text.split('.')[0]
                break

        # 创建一个新的段落对象用于存放表序号题注
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐

        # 创建一个新的表序号题注对象
        caption = OxmlElement('w:fldSimple')
        caption.set(qn('w:instr'), f'SEQ Table \\* ARABIC')
        run = p.add_run()
        run.text = f"\n表 {section_number}-{table_count + 1} {caption_text}"
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        p._p.append(caption)

    except Exception as e:
        print(f"An error occurred in add_table_caption: {e}")


def word_docxinitial(doc, text_font_type='宋体', text_font_size=12, text_font_line_spacing=30,
                     text_header='Header', text_footer='Footer'):
    """
    设置文档的全局字体样式、大小和页眉页脚。

    :param doc: Document对象，要设置样式的Word文档对象。
    :param text_font_type: str，可选，正文的字体类型，默认为'宋体'。
    :param text_font_size: int，可选，正文的字体大小，默认为12磅。
    :param text_font_line_spacing: int，可选，正文的行间距，默认为30磅。
    :param text_header: str，可选，页眉内容，默认为'Header'。
    :param text_footer: str，可选，页脚内容，默认为'Footer'。

    :return: Document对象，设置完样式后的Word文档对象。
    """
    # 设置正文字体类型、大小和行间距
    doc.styles["Normal"].font.name = text_font_type
    doc.styles["Normal"].font.size = Pt(text_font_size)
    doc.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), text_font_type)
    doc.styles["Normal"]._element.line_spacing = Pt(text_font_line_spacing)  # 设置行间距为1.5倍

    # 设置页眉
    header = doc.sections[0].header
    pheader = header.paragraphs[0]  # 获取页眉的第一个段落
    ph_header = pheader.add_run(text_header)
    ph_header.font.name = text_font_type  # 设置页眉字体样式
    ph_header._element.rPr.rFonts.set(qn('w:eastAsia'), text_font_type)
    ph_header.font.size = Pt(text_font_size)  # 设置页眉字体大小
    pheader.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 页眉对齐方式设为居中

    # 设置页脚
    footer = doc.sections[0].footer
    pfooter = footer.paragraphs[0]  # 获取页脚的第一个段落
    ph_footer = pfooter.add_run(text_footer)
    ph_footer.font.name = text_font_type  # 设置页脚字体样式
    ph_footer._element.rPr.rFonts.set(qn('w:eastAsia'), text_font_type)
    ph_footer.font.size = Pt(text_font_size)  # 设置页脚字体大小
    pfooter.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 页眉对齐方式设为居中

    # 返回文档对象
    return doc


def word_get_section_number(paragraph_text):
    """
    从段落文本中提取章节号
    """
    try:
        parts = paragraph_text.split()
        for part in parts:
            if part.replace('.', '').isdigit():
                return part
        return None
    except Exception as e:
        print(f"An error occurred in get_section_number: {e}")
        return None


def word_get_table_cell_text(table, row_index, column_index):
    """
    获取表格中指定单元格的文本值。
    :param table: 表格对象。
    :param row_index: 行索引（0开始）。
    :param column_index: 列索引（0开始）。
    :return: 单元格中的文本值。
    """
    try:
        return table.cell(row_index, column_index).text
    except AttributeError as e:
        print(f"AttributeError: {e} occurred in get_table_cell_text")
        return None
    except IndexError as e:
        print(f"IndexError: {e} occurred in get_table_cell_text")
        return None


def word_merge_cells(table, start_row, start_column, end_row, end_column):
    """
    合并表格中指定范围内的单元格。

    参数：
    table: 表格对象
    start_row: 起始行索引
    start_column: 起始列索引
    end_row: 结束行索引
    end_column: 结束列索引
    """
    # 检查起始行和列是否在有效范围内
    if start_row < 0 or start_row >= len(table.rows):
        raise ValueError("Start row index out of range occurred in merge_cells.")
    if start_column < 0 or start_column >= len(table.columns):
        raise ValueError("Start column index out of range occurred in merge_cells.")

    # 检查结束行和列是否在有效范围内
    if end_row < 0 or end_row >= len(table.rows):
        raise ValueError("End row index out of range occurred in merge_cells.")
    if end_column < 0 or end_column >= len(table.columns):
        raise ValueError("End column index out of range occurred in merge_cells.")

    # 合并单元格
    table.cell(start_row, start_column).merge(table.cell(end_row, end_column))


def word_merge_cells_by_column(table, column_index):
    """
    在表格中根据指定列合并相同值的行，并在合并后的单元格中保留唯一值
    :param table: 表格对象
    :param column_index: 要合并的列索引
    """
    try:
        prev_value = None
        merge_start = None

        for row_index, row in enumerate(table.rows):
            current_cell_value = row.cells[column_index].text.strip()

            if current_cell_value != prev_value:
                if merge_start is not None:
                    # Merge cells from merge_start to current row - 1 in the specified column
                    for i in range(merge_start, row_index):
                        # Keep only the value of the first row in the merged cells
                        table.cell(i, column_index).text = prev_value
                        table.cell(i, column_index).merge(table.cell(row_index - 1, column_index))

                # Update merge_start to the current row
                merge_start = row_index

            # Update prev_value for the next iteration
            prev_value = current_cell_value

        # Merge the last set of cells if needed
        if merge_start is not None and merge_start != len(table.rows):
            for i in range(merge_start, len(table.rows)):
                # Keep only the value of the first row in the merged cells
                table.cell(i, column_index).text = prev_value
                table.cell(i, column_index).merge(table.cell(len(table.rows) - 1, column_index))

    except Exception as e:
        print(f"An error occurred in merge_cells_by_column: {e}")


def word_set_number_to_roman(number):
    """
    将整数转换为罗马数字。
    """
    try:
        roman_numerals = {
            1: 'I', 2: 'II', 3: 'III', 4: 'IV', 5: 'V', 6: 'VI', 7: 'VII', 8: 'VIII', 9: 'IX', 10: 'X',
            20: 'XX', 30: 'XXX', 40: 'XL', 50: 'L', 60: 'LX', 70: 'LXX', 80: 'LXXX', 90: 'XC', 100: 'C',
            200: 'CC', 300: 'CCC', 400: 'CD', 500: 'D', 600: 'DC', 700: 'DCC', 800: 'DCCC', 900: 'CM', 1000: 'M'
        }
        if n in roman_numerals:
            return roman_numerals[n]
        else:
            return ''.join([roman_numerals[int(digit) * 10 ** i] for i, digit in enumerate(reversed(str(n)))][::-1])
    except Exception as e:
        print(f"An error occurred in set_number_to_roman: {e}")


def word_set_table_cell_text(table, row_index, column_index, text):
    """
    设置表格中指定单元格的文本内容。

    参数：
    table: 表格对象
    row_index: 行索引（0开始）
    column_index: 列索引（0开始）
    text: 要设置的文本内容
    """
    try:
        table.cell(row_index, column_index).text = text
    except Exception as e:
        print(f"An error occurred in set_table_cell_text: {e}")


def word_set_table_alignment(table, start_row=None, start_column=None, end_row=None, end_column=None,
                             horizontal_alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             vertical_alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """
    设置表格中指定范围内的单元格的文本对齐方式，默认情况下将所有单元格的文本设置为水平和垂直居中对齐。

    参数：
    table: 表格对象
    start_row: 起始行索引，默认为 None（默认第一行）
    start_column: 起始列索引，默认为 None（默认第一列）
    end_row: 结束行索引，默认为 None（默认最后一行）
    end_column: 结束列索引，默认为 None（默认最后一列）
    horizontal_alignment: 水平对齐方式，默认为水平居中对齐
    vertical_alignment: 垂直对齐方式，默认为垂直居中对齐
    """
    try:
        # 默认情况下，如果未提供范围，则将范围设置为整个表格
        if start_row is None:
            start_row = 0
        if start_column is None:
            start_column = 0
        if end_row is None:
            end_row = len(table.rows) - 1
        if end_column is None:
            end_column = len(table.columns) - 1

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if (start_row <= i <= end_row and start_column <= j <= end_column) or (
                        start_row > end_row and start_column > end_column):
                    cell.paragraphs[0].alignment = horizontal_alignment
                    cell.vertical_alignment = vertical_alignment

    except Exception as e:
        print(f"An error occurred in set_table_alignment: {e}")


def word_set_table_cell_text(table, row_index, column_index, text):
    """
    设置表格中指定单元格的文本内容。

    参数：
    table: 表格对象
    row_index: 行索引（0开始）
    column_index: 列索引（0开始）
    text: 要设置的文本内容
    """
    try:
        table.cell(row_index, column_index).text = text
    except Exception as e:
        print(f"An error occurred in set_table_cell_text: {e}")


def word_set_table_style(table, font_size=10.5, spacing_before=0):
    """
    设置表格的字体大小和段前间距。
    :param table: 表格对象。
    :param font_size: 字体大小，默认为10.5。
    :param spacing_before: 段前间距，默认为0。
    :return:
    """
    try:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)  # 设置字体大小
                    paragraph.paragraph_format.space_before = Pt(spacing_before)  # 设置段前间距
    except Exception as e:
        print(f"An error occurred in set_table_style: {e}")


def word_setsectionformat(doc):
    """
    在doc结尾处调用进行设置
    通过sections（节）进行设置时，节对应文档中的每一页,每个节在没输入内容之前是不存在的,因此在最后才对每个节逐一进行设置
    :param doc:
    :return:
    """
    for sec in doc.sections:
        # 设置页面边距(左上25毫米，右下15毫米)
        sec.top_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        # 设置纸张大小(A4)
        sec.page_height = Mm(297)
        sec.page_width = Mm(210)
        # 设置页眉页脚距离
        sec.header_distance = Cm(1.5)
        sec.footer_distance = Cm(0.2)
